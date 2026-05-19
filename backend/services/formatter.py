import os
import re
import PyPDF2
from docx import Document
from services.equation_formatter import EquationFormatter
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import shutil

def hex_to_rgb(hex_color):
    """Convert hex color string to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    if len(hex_color) == 3:
        hex_color = ''.join([c*2 for c in hex_color])
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def extract_text(filepath):
    """Extract text from PDF or DOCX file."""
    text = ""
    if filepath.lower().endswith('.pdf'):
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    elif filepath.lower().endswith('.docx'):
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
    return text

def parse_text(text):
    """Heuristic-based parser to identify sections of the paper."""
    lines = text.split('\n')
    lines = [L.strip() for L in lines if L.strip()]
    
    parsed = {
        'title': '',
        'authors': [],
        'university': '',
        'abstract': '',
        'keywords': '',
        'body': [],
        'references': []
    }
    
    if not lines:
        return parsed

    state = 'start'
    
    for line in lines:
        lower_line = line.lower()
        
        if state == 'start':
            parsed['title'] = line
            state = 'authors'
        elif state == 'authors':
            if lower_line.startswith('abstract'):
                parsed['abstract'] += line[8:].strip() + " "
                state = 'abstract'
            elif 'university' in lower_line or 'department' in lower_line or 'college' in lower_line or 'institute' in lower_line:
                parsed['university'] = line
            else:
                if len(line.split()) < 10 and not parsed['abstract']:
                    parsed['authors'].append(line)
                else:
                    # Maybe it's abstract without a header
                    if len(line.split()) > 20:
                        parsed['abstract'] += line + " "
                        state = 'abstract'
        elif state == 'abstract':
            if lower_line.startswith('keyword') or lower_line.startswith('index term'):
                parsed['keywords'] = line
                state = 'body'
            elif lower_line.startswith('introduction') or (len(line.split()) < 10 and not line.endswith('.')):
                parsed['body'].append(line)
                state = 'body'
            else:
                parsed['abstract'] += line + " "
        elif state == 'body':
            if lower_line == 'references' or lower_line == 'works cited' or lower_line == 'bibliography':
                state = 'references'
            else:
                parsed['body'].append(line)
        elif state == 'references':
            # Collect references
            if line:
                parsed['references'].append(line)
                
    return parsed

def set_number_of_columns(section, cols, space_twips):
    sectPr = section._sectPr
    cols_element = sectPr.xpath('./w:cols')
    if not cols_element:
        cols_element = OxmlElement('w:cols')
        sectPr.append(cols_element)
    else:
        cols_element = cols_element[0]
    cols_element.set(qn('w:num'), str(cols))
    cols_element.set(qn('w:space'), str(space_twips))

def format_document(parsed_data, output_path, options=None):
    """Create a formatted docx document based on parsed data and academic guidelines."""
    if options is None:
        options = {}
    
    citation_style = options.get('citation_style', 'APA')
    num_columns = int(options.get('columns', 2))
    
    h_font_name = options.get('heading_font', 'Times New Roman')
    h_font_size = int(options.get('heading_size', 20))
    h_color_rgb = hex_to_rgb(options.get('heading_color', '#000000'))
    
    c_font_name = options.get('content_font', 'Times New Roman')
    c_font_size = int(options.get('content_size', 10))
    c_color_rgb = hex_to_rgb(options.get('content_color', '#000000'))

    doc = Document()
    
    # Configure default style
    style = doc.styles['Normal']
    font = style.font
    font.name = c_font_name
    font.size = Pt(c_font_size)
    font.color.rgb = RGBColor(*c_color_rgb)
    style.paragraph_format.line_spacing = 1.0  # Single-spaced (typical in IEEE/IJRTI)
    
    # 0.6-inch margins for all sections
    # Initial section (Title, Authors, Abstract) -> 1 column
    section1 = doc.sections[0]
    section1.top_margin = Inches(0.6)
    section1.bottom_margin = Inches(0.6)
    section1.left_margin = Inches(0.6)
    section1.right_margin = Inches(0.6)
    
    # Title (Bold, Centered, 20pt)
    if parsed_data['title']:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(parsed_data['title'])
        title_run.bold = True
        title_run.font.name = h_font_name
        title_run.font.size = Pt(h_font_size)
        title_run.font.color.rgb = RGBColor(*h_color_rgb)
        
    # Authors
    for author in parsed_data['authors']:
        auth_p = doc.add_paragraph()
        auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auth_run = auth_p.add_run(author)
        auth_run.bold = True
        auth_run.font.size = Pt(11)
        
    # University
    if parsed_data['university']:
        univ_p = doc.add_paragraph()
        univ_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        u_run = univ_p.add_run(parsed_data['university'])
        u_run.font.size = Pt(10)
        
    # Abstract
    abs_text = parsed_data['abstract'].strip()
    if abs_text:
        abs_p = doc.add_paragraph()
        abs_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        a_title = abs_p.add_run("Abstract— ")
        a_title.bold = True
        a_title.italic = True
        a_body = abs_p.add_run(abs_text)
        a_body.bold = True
        a_body.italic = True
        
    # Keywords
    if parsed_data['keywords']:
        kw_p = doc.add_paragraph()
        kw_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        kw_run = kw_p.add_run(parsed_data['keywords'])
        kw_run.bold = True
        kw_run.italic = True
        
    # Create Continuous Section for two columns
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.top_margin = Inches(0.6)
    new_section.bottom_margin = Inches(0.6)
    new_section.left_margin = Inches(0.6)
    new_section.right_margin = Inches(0.6)
    # columns with 0.3 inch spacing (432 twips)
    set_number_of_columns(new_section, num_columns, 432)
        
    # Body
    for para in parsed_data['body']:
        words = para.split()
        if len(words) < 12 and not para.endswith('.'):
            # Heading
            h_p = doc.add_paragraph()
            h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h_run = h_p.add_run(para)
            h_run.bold = True
            h_run.font.name = h_font_name
            h_run.font.size = Pt(int(h_font_size * 0.6)) # Smaller sub-heading
            h_run.font.color.rgb = RGBColor(*h_color_rgb)
        else:
            # Regular paragraph
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.15)
            
    # References
    if parsed_data['references']:
        ref_h = doc.add_paragraph()
        ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_h.add_run("References").bold = True
        
        for line in parsed_data['references']:
            r_p = doc.add_paragraph(line)
            r_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_p.paragraph_format.left_indent = Inches(0.15)
            r_p.paragraph_format.first_line_indent = Inches(-0.15)

    try:
        doc.save(output_path)
        return True, "Formatting successful."
    except Exception as e:
        return False, str(e)

def validate_academic_paper(text):
    """
    Validates if the text content matches a standard research paper.
    Checks for the presence of Abstract, Introduction/Methodology, and References.
    """
    lower_text = text.lower()
    
    # 1. Check for abstract
    has_abstract = "abstract" in lower_text
    
    # 2. Check for references
    has_references = any(ref_kw in lower_text for ref_kw in ["references", "works cited", "bibliography"])
    
    # 3. Check for typical academic structural headers
    has_introduction = "introduction" in lower_text
    
    if not has_abstract:
        return False, "This document does not contain an 'Abstract' section. A valid academic research paper must start with an Abstract summarizing the study's purpose and findings."
        
    if not has_references:
        return False, "This document does not contain a 'References' or 'Bibliography' section. Legitimate academic papers must cite sources at the end."
        
    if not has_introduction:
        return False, "This document does not contain an 'Introduction' section. Academic research papers must have a clear structural introduction."
        
    return True, ""

def in_place_format_docx(input_path, output_path, options=None):
    """Format docx in place to retain tables, images, and inline formatting"""
    if options is None:
        options = {}
        
    num_columns = int(options.get('columns', 2))
    
    h_font_name = options.get('heading_font', 'Times New Roman')
    h_font_size = int(options.get('heading_size', 20))
    h_color_rgb = hex_to_rgb(options.get('heading_color', '#000000'))
    
    c_font_name = options.get('content_font', 'Times New Roman')
    c_font_size = int(options.get('content_size', 10))
    c_color_rgb = hex_to_rgb(options.get('content_color', '#000000'))

    try:
        shutil.copy(input_path, output_path)
        doc = Document(output_path)
    except Exception as e:
        return False, f"Could not create output file: {str(e)}"

    # Configure default style
    style = doc.styles['Normal']
    font = style.font
    font.name = c_font_name
    font.size = Pt(c_font_size)
    font.color.rgb = RGBColor(*c_color_rgb)
    if hasattr(style.paragraph_format, 'line_spacing'):
        style.paragraph_format.line_spacing = 1.0

    # 1. Identify where body starts (e.g. Introduction or after keywords)
    body_start_idx = None
    ref_start_idx = None
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        text_lower = text.lower()
        if not text:
            continue
            
        # Detect where body starts
        if body_start_idx is None:
            if text_lower.startswith("introduction") or text_lower.startswith("i. introduction") or re.match(r'^i\.\s+', text_lower):
                body_start_idx = i
                
        # Detect where references start
        if text_lower in ['references', 'works cited', 'bibliography'] or text_lower.startswith('references') or text_lower.startswith('bibliography'):
            ref_start_idx = i

    # Fallback for body start
    if body_start_idx is None:
        for i, p in enumerate(doc.paragraphs):
            text_lower = p.text.strip().lower()
            if "keyword" in text_lower or "index terms" in text_lower:
                body_start_idx = i + 1
                break
        if body_start_idx is None:
            body_start_idx = min(5, len(doc.paragraphs) - 1)

    # 2. Inject Continuous Section break right before body starts
    if body_start_idx is not None and body_start_idx > 0:
        prev_p = doc.paragraphs[body_start_idx - 1]
        pPr = prev_p._element.get_or_add_pPr()
        sectPr = pPr.xpath('w:sectPr')
        if not sectPr:
            sectPr_el = OxmlElement('w:sectPr')
            type_el = OxmlElement('w:type')
            type_el.set(qn('w:val'), 'continuous')
            sectPr_el.append(type_el)
            pPr.append(sectPr_el)

    # 3. Configure column count and margins on all sections
    # Refresh doc reference to reload newly injected sections
    try:
        doc.save(output_path)
        doc = Document(output_path)
    except:
        pass
    
    for idx, sec in enumerate(doc.sections):
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.6)
        sec.right_margin = Inches(0.6)
        
        if idx == 0:
            # Title & Abstract section spans full width (1 column)
            set_number_of_columns(sec, 1, 0)
        else:
            # Body section is 2 columns (or user options)
            set_number_of_columns(sec, num_columns, 432) # 0.3 in column spacing

    # 4. Standardize headings and body paragraphs formatting
    title_found = False
    in_references = False
    
    heading_map = {
        'introduction': 'I. INTRODUCTION',
        'related work': 'II. RELATED WORK',
        'literature review': 'II. LITERATURE REVIEW',
        'methodology': 'III. METHODOLOGY',
        'methods': 'III. METHODOLOGY',
        'results': 'IV. RESULTS',
        'discussion': 'V. DISCUSSION',
        'conclusion': 'VI. CONCLUSION',
        'references': 'REFERENCES'
    }

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue
            
        text_lower = text.lower()
        words = text.split()
        num_words = len(words)
        
        # Heading styles or patterns
        is_heading_style = p.style.name.startswith('Heading') or p.style.name == 'Title'
        text_is_title_case = text.istitle() and num_words > 1
        runs_are_bold = sum(1 for r in p.runs if r.bold) > len(p.runs) / 2 if p.runs else False
        
        # Check if we transitioned to references
        if ref_start_idx is not None and i >= ref_start_idx:
            in_references = True

        # Heading detection and Roman numeral formatting
        matched_heading_key = None
        for key in heading_map.keys():
            if text_lower == key or text_lower.endswith(" " + key) or text_lower == heading_map[key].lower():
                matched_heading_key = key
                break
                
        if not title_found and i < 4 and num_words < 20:
            # Title centering and formatting
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            for run in p.runs:
                run.font.name = h_font_name
                run.font.size = Pt(h_font_size)
                run.font.color.rgb = RGBColor(*h_color_rgb)
                run.bold = True
            if num_words > 3:
                title_found = True
        elif text_lower.startswith('abstract') or text_lower == 'abstract':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            # IEEE style abstract format
            for run in p.runs:
                run.font.name = h_font_name
                run.font.size = Pt(c_font_size)
                run.font.color.rgb = RGBColor(*c_color_rgb)
                run.bold = True
                run.italic = True
        elif matched_heading_key is not None or is_heading_style or (num_words < 12 and not text.endswith('.') and (text_is_title_case or runs_are_bold)):
            # Standardized Heading
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            heading_text = heading_map[matched_heading_key] if matched_heading_key else text.upper()
            
            # Reset text with proper Roman numerals
            p.text = heading_text
            for run in p.runs:
                run.font.name = h_font_name
                run.font.size = Pt(int(h_font_size * 0.55))
                run.font.color.rgb = RGBColor(*h_color_rgb)
                run.bold = True
        elif in_references:
            # References format: Justified, hanging indent
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            for run in p.runs:
                run.font.name = c_font_name
                run.font.size = Pt(c_font_size)
                run.font.color.rgb = RGBColor(*c_color_rgb)
        else:
            # Regular body paragraph
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.first_line_indent = Inches(0.15)
            
            # Auto-format LaTeX formulas first
            has_equations = EquationFormatter.process_paragraph_equations(p)
            
            if not has_equations:
                for run in p.runs:
                    run.font.name = c_font_name
                    run.font.size = Pt(c_font_size)
                    run.font.color.rgb = RGBColor(*c_color_rgb)

    # Center paragraphs containing images
    for p in doc.paragraphs:
        if p._element.xpath('.//pic:pic') or p._element.xpath('.//w:drawing'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Inches(0)
            
    # Remove trailing empty paragraphs via direct body child traversal (DOM accurate)
    body = doc._element.body
    for child in reversed(list(body)):
        if child.tag.endswith('sectPr'):
            continue
        if child.tag.endswith('p'):
            p = child
            text = "".join(node.text for node in p.iter() if node.tag.endswith('t') and node.text)
            has_drawing = bool(p.xpath('.//w:drawing') or p.xpath('.//pic:pic'))
            if not text.strip() and not has_drawing:
                body.remove(child)
            else:
                break
        else:
            break
                    
    # Autofit tables to ensure they don't bleed out of columns
    for table in doc.tables:
        table.autofit = True
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            tblW = tblPr[0].xpath('w:tblW')
            if not tblW:
                tblW_el = OxmlElement('w:tblW')
                tblPr[0].append(tblW_el)
            else:
                tblW_el = tblW[0]
            tblW_el.set(qn('w:type'), 'pct')
            tblW_el.set(qn('w:w'), '5000')  # 100% of current column width
            
        # Strip cell width constraints so table can actually shrink organically
        for row in table.rows:
            for cell in row.cells:
                tcW = cell._element.xpath('.//w:tcW')
                for w in tcW:
                    w.set(qn('w:type'), 'auto')
                    w.set(qn('w:w'), '0')
                # Prevent cells from accidentally page breaking layout
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = False
                    p.paragraph_format.page_break_before = False
            
    # Autofit inline images to ensure they don't break columns
    MAX_WIDTH = Inches(3.25)
    for shape in doc.inline_shapes:
        if shape.width > MAX_WIDTH:
            ratio = MAX_WIDTH / shape.width
            shape.width = MAX_WIDTH
            shape.height = int(shape.height * ratio)

    try:
        doc.save(output_path)
        return True, "Formatting successful."
    except Exception as e:
        return False, str(e)

def mock_ai_refinement(text, instruction):
    """Simulate AI text refinement based on instructions."""
    inst = instruction.lower()
    
    # 1. Formalize / Academic
    if "formal" in inst or "academic" in inst or "formalize" in inst:
        academic_replacements = {
            "get": "obtain",
            "don't": "do not",
            "it's": "it is",
            "can't": "cannot",
            "a lot of": "numerous",
            "show": "demonstrate",
            "find out": "determine",
            "good": "advantageous",
            "bad": "detrimental"
        }
        refined = text
        for old, new in academic_replacements.items():
            refined = refined.replace(old, new).replace(old.capitalize(), new.capitalize())
        return refined
        
    # 2. Shorten
    elif "shorten" in inst or "concise" in inst:
        points = text.split(". ")
        if len(points) > 1:
            return f"{points[0]}. (Summary: {points[-1]})"
        return f"{text[:len(text)//2]}..."
        
    # 3. Explain / Simple
    elif "explain" in inst or "simpler" in inst:
        return f"[Simpler Explanation]: Essentially, this means that {text.lower() if text[0].islower() else text[0].lower() + text[1:]}"
        
    # 4. Grammar
    elif "grammar" in inst or "punctuation" in inst:
        # Simulated grammar fix (e.g. capitalize first letter, add period)
        refined = text.strip()
        if refined and not refined[0].isupper():
            refined = refined[0].upper() + refined[1:]
        if refined and not refined.endswith("."):
            refined += "."
        return refined
        
    # 5. Default/Custom
    return f"[AI-Edited]: {text}"

def refine_docx(file_path, original_text, instruction):
    """Find text in docx and replace it with refined version."""
    try:
        doc = Document(file_path)
        new_text = mock_ai_refinement(original_text, instruction)
        found = False
        
        # Search and replace logic
        for para in doc.paragraphs:
            if original_text in para.text:
                # Replace in runs to preserve some formatting if possible
                # Simple full paragraph replace for now to ensure consistency
                para.text = para.text.replace(original_text, new_text)
                found = True
                
        if not found:
            # Try fuzzy matching if exact fails
            for para in doc.paragraphs:
                if len(para.text) > 10 and original_text[:10] in para.text:
                    para.text = para.text.replace(para.text, new_text)
                    found = True
                    break
        
        doc.save(file_path)
        return True, new_text
    except Exception as e:
        return False, str(e)
