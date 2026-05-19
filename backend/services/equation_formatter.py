import re
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

class EquationFormatter:
    """
    Parses LaTeX and standard equations from paragraphs and translates them
    into Office Math Markup Language (OMML) XML structures that render natively in MS Word.
    """
    
    # Standard OMML Math namespace declarations for parsing
    OMML_NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    
    @staticmethod
    def process_paragraph_equations(p) -> bool:
        """
        Scans a python-docx Paragraph for LaTeX equations (e.g. $e=mc^2$ or $$E=mc^2$$)
        and replaces the raw text run with native Office Math elements.
        Returns True if any equations were processed.
        """
        text = p.text
        if not text:
            return False
            
        # Detect inline equations ($...$) and block equations ($$...$$)
        # Avoid matching double dollar signs as single dollar signs by checking block first
        block_pattern = r'\$\$(.*?)\$\$'
        inline_pattern = r'\$(.*?)\$'
        
        has_block = re.search(block_pattern, text)
        has_inline = re.search(inline_pattern, text)
        
        if not has_block and not has_inline:
            return False
            
        # Clear existing runs in the paragraph and rebuild them semantically
        # First, extract raw text segments and equation blocks
        p.text = "" # Clears paragraph
        
        # Simple parser to segment text and equations sequentially
        current_idx = 0
        matches = []
        
        # Combine block and inline equations sorted by their appearance index
        for m in re.finditer(block_pattern, text):
            matches.append((m.start(), m.end(), m.group(1), True))
        for m in re.finditer(inline_pattern, text):
            # Ensure we don't double match block equations
            start, end = m.start(), m.end()
            is_inside_block = any(b_start <= start < b_end for b_start, b_end, _, _ in matches if _ == True)
            if not is_inside_block:
                matches.append((start, end, m.group(1), False))
                
        matches.sort(key=lambda x: x[0])
        
        if not matches:
            p.add_run(text)
            return False
            
        # Build paragraph runs step by step
        for start, end, equation_text, is_block in matches:
            # 1. Add preceding normal text run
            if start > current_idx:
                p.add_run(text[current_idx:start])
                
            # 2. Translate LaTeX to OMML XML
            omml_xml = EquationFormatter.latex_to_omml(equation_text, is_block)
            
            # 3. Insert XML element directly into paragraph's DOM element
            try:
                math_element = parse_xml(omml_xml)
                p._element.append(math_element)
            except Exception as e:
                # Fallback to plain text run if XML parse fails
                fallback_run = p.add_run(f" {equation_text} ")
                fallback_run.italic = True
                print(f"Failed parsing OMML equation XML: {e}")
                
            current_idx = end
            
        # Add remaining trailing text
        if current_idx < len(text):
            p.add_run(text[current_idx:])
            
        return True

    @staticmethod
    def latex_to_omml(latex: str, is_block: bool = False) -> str:
        """
        Translates raw LaTeX notation into a valid MS Word OMML XML string.
        Supports: exponents (^), subscripts (_), fractions (\\frac), greek letters, operators, roots.
        """
        latex = latex.strip()
        
        # Initialize math structural elements
        body_xml = EquationFormatter._parse_latex_recursive(latex)
        
        if is_block:
            # Block equations are wrapped in a Math Paragraph container (m:oMathPara) for full spacing/centering
            return f'<m:oMathPara {EquationFormatter.OMML_NS}><m:oMath>{body_xml}</m:oMath></m:oMathPara>'
        else:
            # Inline equations sit inline with regular text paragraphs
            return f'<m:oMath {EquationFormatter.OMML_NS}>{body_xml}</m:oMath>'

    @staticmethod
    def _parse_latex_recursive(latex: str) -> str:
        """
        Recursively translates LaTeX constructs (fractions, subscripts, superscripts, etc.) into OMML blocks.
        """
        if not latex:
            return ""
            
        # Match high-priority constructs first: Fractions \frac{num}{den}
        frac_match = re.search(r'\\frac\s*{(.*?)}{(.*?)}', latex)
        if frac_match:
            num = EquationFormatter._parse_latex_recursive(frac_match.group(1))
            den = EquationFormatter._parse_latex_recursive(frac_match.group(2))
            
            # OMML Fraction tag structure
            frac_xml = f'<m:f><m:num><m:r><m:t>{num}</m:t></m:r></m:num><m:den><m:r><m:t>{den}</m:t></m:r></m:den></m:f>'
            
            # Replace matching block and parse surrounding text recursively
            before = EquationFormatter._parse_latex_recursive(latex[:frac_match.start()])
            after = EquationFormatter._parse_latex_recursive(latex[frac_match.end():])
            return before + frac_xml + after

        # Exponent rules (base^power)
        power_match = re.search(r'([a-zA-Z0-9\(\)]+)\^([a-zA-Z0-9\+\-\=]+|\{[a-zA-Z0-9\+\-\=]+\})', latex)
        if power_match:
            base = power_match.group(1)
            power = power_match.group(2).strip('{}')
            
            base_parsed = EquationFormatter._parse_latex_recursive(base)
            power_parsed = EquationFormatter._parse_latex_recursive(power)
            
            # OMML Superscript tag structure
            power_xml = f'<m:sSup><m:e><m:r><m:t>{base_parsed}</m:t></m:r></m:e><m:sup><m:r><m:t>{power_parsed}</m:t></m:r></m:sup></m:sSup>'
            
            before = EquationFormatter._parse_latex_recursive(latex[:power_match.start()])
            after = EquationFormatter._parse_latex_recursive(latex[power_match.end():])
            return before + power_xml + after

        # Subscript rules (base_subscript)
        sub_match = re.search(r'([a-zA-Z0-9\(\)]+)_([a-zA-Z0-9\+\-\=]+|\{[a-zA-Z0-9\+\-\=]+\})', latex)
        if sub_match:
            base = sub_match.group(1)
            sub = sub_match.group(2).strip('{}')
            
            base_parsed = EquationFormatter._parse_latex_recursive(base)
            sub_parsed = EquationFormatter._parse_latex_recursive(sub)
            
            # OMML Subscript tag structure
            sub_xml = f'<m:sSub><m:e><m:r><m:t>{base_parsed}</m:t></m:r></m:e><m:sub><m:r><m:t>{sub_parsed}</m:t></m:r></m:sub></m:sSub>'
            
            before = EquationFormatter._parse_latex_recursive(latex[:sub_match.start()])
            after = EquationFormatter._parse_latex_recursive(latex[sub_match.end():])
            return before + sub_xml + after
            
        # Replace Greek letters with actual math symbols
        greek_symbols = {
            r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ', r'\epsilon': 'ε',
            r'\zeta': 'ζ', r'\eta': 'η', r'\theta': 'θ', r'\iota': 'ι', r'\kappa': 'κ',
            r'\lambda': 'λ', r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ', r'\pi': 'π',
            r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ', r'\upsilon': 'υ', r'\phi': 'φ',
            r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω', r'\Sigma': 'Σ', r'\Delta': 'Δ',
            r'\theta': 'Θ', r'\lambda': 'Λ', r'\omega': 'Ω', r'\infty': '∞', r'\sum': '∑',
            r'\int': '∫', r'\approx': '≈', r'\ne': '≠', r'\le': '≤', r'\ge': '≥'
        }
        
        working_latex = latex
        for latex_sym, unicode_sym in greek_symbols.items():
            working_latex = working_latex.replace(latex_sym, unicode_sym)
            
        # Return as basic OMML run (r) holding math text (t)
        return f'<m:r><m:t>{working_latex}</m:t></m:r>'
